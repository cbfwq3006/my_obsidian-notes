from __future__ import annotations

import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


SUPPORTED = {".docx", ".xlsx"}
UNSUPPORTED = {".doc", ".xls", ".wps", ".jpg", ".jpeg", ".png", ".rar", ".pg"}


def safe_part(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name)
    return name.rstrip(" .") or "_"


def relative_output_path(source: Path, root: Path, out_root: Path) -> Path:
    rel = source.relative_to(root)
    safe_parts = [safe_part(part) for part in rel.parts]
    target = out_root.joinpath(*safe_parts)
    return target.with_suffix(".md")


def escape_md_cell(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("|", "\\|").replace("\n", "<br>")
    return text.strip()


def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield "paragraph", child
        elif child.tag.endswith("}tbl"):
            yield "table", child


def paragraph_text_from_element(document: Document, element) -> str:
    for paragraph in document.paragraphs:
        if paragraph._p is element:
            return paragraph.text.strip()
    return ""


def table_from_element(document: Document, element):
    for table in document.tables:
        if table._tbl is element:
            return table
    return None


def table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(escape_md_cell(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(escape_md_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def convert_docx(source: Path, target: Path) -> dict:
    doc = Document(str(source))
    lines: list[str] = [
        f"# {source.stem}",
        "",
        f"Source: `{source}`",
        "",
    ]
    paragraph_count = 0
    table_count = 0

    for kind, element in iter_block_items(doc):
        if kind == "paragraph":
            text = paragraph_text_from_element(doc, element)
            if text:
                paragraph_count += 1
                lines.append(text)
                lines.append("")
        elif kind == "table":
            table = table_from_element(doc, element)
            if table is None:
                continue
            table_count += 1
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            md_table = table_to_markdown(rows)
            if md_table:
                lines.append(md_table)
                lines.append("")

    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return {
        "paragraphs": paragraph_count,
        "tables": table_count,
    }


def sheet_to_markdown(ws) -> str:
    rows = list(ws.iter_rows(values_only=True))
    while rows and all(value is None for value in rows[-1]):
        rows.pop()
    if not rows:
        return "_Empty sheet._"

    max_col = max((len(row) for row in rows), default=0)
    compact_rows = []
    for row in rows:
        values = list(row) + [None] * (max_col - len(row))
        compact_rows.append([escape_md_cell(value) for value in values])

    return table_to_markdown(compact_rows)


def convert_xlsx(source: Path, target: Path) -> dict:
    wb = load_workbook(filename=str(source), data_only=True, read_only=True)
    lines: list[str] = [
        f"# {source.stem}",
        "",
        f"Source: `{source}`",
        "",
    ]
    sheet_count = 0
    for ws in wb.worksheets:
        sheet_count += 1
        lines.append(f"## Sheet: {ws.title}")
        lines.append("")
        lines.append(sheet_to_markdown(ws))
        lines.append("")

    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return {
        "sheets": sheet_count,
    }


def convert_file(source: Path, root: Path, out_root: Path) -> dict:
    target = relative_output_path(source, root, out_root)
    ext = source.suffix.lower()
    if ext == ".docx":
        stats = convert_docx(source, target)
    elif ext == ".xlsx":
        stats = convert_xlsx(source, target)
    else:
        raise ValueError(f"Unsupported extension: {ext}")

    return {
        "source": str(source),
        "target": str(target),
        "extension": ext,
        "status": "converted",
        **stats,
    }


def write_csv(path: Path, rows: list[dict], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    source_root = Path(args.source).resolve()
    output_root = Path(args.output).resolve()
    output_root.mkdir(parents=True, exist_ok=True)

    converted: list[dict] = []
    failed: list[dict] = []
    unsupported: list[dict] = []

    for source in sorted(source_root.rglob("*")):
        if not source.is_file():
            continue
        ext = source.suffix.lower()
        if ext in SUPPORTED:
            try:
                converted.append(convert_file(source, source_root, output_root))
            except Exception as exc:
                failed.append({
                    "source": str(source),
                    "extension": ext,
                    "status": "failed",
                    "error": repr(exc),
                })
        else:
            unsupported.append({
                "source": str(source),
                "extension": ext or "(none)",
                "status": "unsupported",
            })

    manifest = {
        "source_root": str(source_root),
        "output_root": str(output_root),
        "converted_count": len(converted),
        "failed_count": len(failed),
        "unsupported_count": len(unsupported),
        "converted": converted,
        "failed": failed,
        "unsupported": unsupported,
    }
    (output_root / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    write_csv(
        output_root / "converted.csv",
        converted,
        ["source", "target", "extension", "status", "paragraphs", "tables", "sheets"],
    )
    write_csv(
        output_root / "failed.csv",
        failed,
        ["source", "extension", "status", "error"],
    )
    write_csv(
        output_root / "unsupported.csv",
        unsupported,
        ["source", "extension", "status"],
    )

    index_lines = [
        "# Conversion Index",
        "",
        f"Source root: `{source_root}`",
        f"Output root: `{output_root}`",
        "",
        f"- Converted: {len(converted)}",
        f"- Failed: {len(failed)}",
        f"- Unsupported: {len(unsupported)}",
        "",
        "## Converted Files",
        "",
    ]
    for row in converted:
        index_lines.append(f"- [{Path(row['target']).name}]({row['target']}) <- `{row['source']}`")

    if failed:
        index_lines.extend(["", "## Failed Files", ""])
        for row in failed:
            index_lines.append(f"- `{row['source']}`: {row['error']}")

    if unsupported:
        index_lines.extend(["", "## Unsupported Files", ""])
        for row in unsupported:
            index_lines.append(f"- `{row['source']}` ({row['extension']})")

    (output_root / "index.md").write_text("\n".join(index_lines) + "\n", encoding="utf-8")

    print(json.dumps({
        "converted": len(converted),
        "failed": len(failed),
        "unsupported": len(unsupported),
        "output": str(output_root),
    }, ensure_ascii=False))
    return 0 if not failed else 1


if __name__ == "__main__":
    raise SystemExit(main())
